"""
문서 로더 모듈

다양한 형식의 문서를 로드하는 클래스를 제공합니다.
"""

import logging
import os
import re
from typing import Dict, List, Optional, Union, Any, Tuple
from pathlib import Path

# 필요한 패키지 임포트
try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    import markdown
except ImportError:
    markdown = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import docx
except ImportError:
    docx = None

try:
    import requests
except ImportError:
    requests = None

from bombay.document_processing.document import Document

logger = logging.getLogger(__name__)

class DocumentLoader:
    """문서 로더 기본 클래스"""
    
    def __init__(self, **kwargs):
        """
        문서 로더 초기화
        
        Args:
            **kwargs: 추가 매개변수
        """
        self.kwargs = kwargs
    
    def load(self, source: str) -> Document:
        """
        문서 로드
        
        Args:
            source: 문서 소스 (파일 경로 등)
            
        Returns:
            Document 객체
        """
        raise NotImplementedError("이 메서드는 하위 클래스에서 구현해야 합니다.")
    
    def _extract_metadata(self, source: str) -> Dict[str, Any]:
        """
        메타데이터 추출
        
        Args:
            source: 문서 소스
            
        Returns:
            메타데이터 딕셔너리
        """
        # 기본 메타데이터
        metadata = {
            "source": source,
            "loader": self.__class__.__name__
        }
        
        # 파일인 경우 추가 메타데이터
        if os.path.isfile(source):
            file_path = Path(source)
            metadata.update({
                "file_path": str(file_path),
                "file_name": file_path.name,
                "file_extension": file_path.suffix.lower()[1:],
                "file_size": os.path.getsize(source),
                "created_at": os.path.getctime(source),
                "modified_at": os.path.getmtime(source)
            })
        
        return metadata


class TextLoader(DocumentLoader):
    """텍스트 파일 로더"""
    
    def __init__(self, encoding: str = "utf-8", **kwargs):
        """
        텍스트 파일 로더 초기화
        
        Args:
            encoding: 파일 인코딩
            **kwargs: 추가 매개변수
        """
        super().__init__(**kwargs)
        self.encoding = encoding
    
    def load(self, source: str) -> Document:
        """
        텍스트 파일 로드
        
        Args:
            source: 파일 경로
            
        Returns:
            Document 객체
        """
        try:
            with open(source, 'r', encoding=self.encoding) as f:
                content = f.read()
            
            metadata = self._extract_metadata(source)
            
            return Document(content=content, metadata=metadata)
        
        except Exception as e:
            logger.error(f"텍스트 파일 로드 중 오류 발생: {e}")
            return Document(content="", metadata={"error": str(e), "source": source})


class MarkdownLoader(DocumentLoader):
    """마크다운 파일 로더"""
    
    def __init__(self, encoding: str = "utf-8", convert_to_html: bool = False, **kwargs):
        """
        마크다운 파일 로더 초기화
        
        Args:
            encoding: 파일 인코딩
            convert_to_html: HTML로 변환 여부
            **kwargs: 추가 매개변수
        """
        super().__init__(**kwargs)
        self.encoding = encoding
        self.convert_to_html = convert_to_html
        
        if convert_to_html and markdown is None:
            logger.warning("markdown 패키지가 설치되지 않았습니다. HTML 변환이 비활성화됩니다.")
            self.convert_to_html = False
    
    def load(self, source: str) -> Document:
        """
        마크다운 파일 로드
        
        Args:
            source: 파일 경로
            
        Returns:
            Document 객체
        """
        try:
            with open(source, 'r', encoding=self.encoding) as f:
                content = f.read()
            
            metadata = self._extract_metadata(source)
            
            if self.convert_to_html and markdown is not None:
                html_content = markdown.markdown(content)
                metadata["content_type"] = "html"
                return Document(content=html_content, metadata=metadata)
            
            return Document(content=content, metadata=metadata)
        
        except Exception as e:
            logger.error(f"마크다운 파일 로드 중 오류 발생: {e}")
            return Document(content="", metadata={"error": str(e), "source": source})


class PDFLoader(DocumentLoader):
    """PDF 파일 로더"""
    
    def __init__(self, extract_images: bool = False, **kwargs):
        """
        PDF 파일 로더 초기화
        
        Args:
            extract_images: 이미지 추출 여부
            **kwargs: 추가 매개변수
        """
        super().__init__(**kwargs)
        self.extract_images = extract_images
        
        if PyPDF2 is None:
            logger.warning("PyPDF2 패키지가 설치되지 않았습니다. PDF 로드가 제한됩니다.")
        
        if extract_images and fitz is None:
            logger.warning("PyMuPDF(fitz) 패키지가 설치되지 않았습니다. 이미지 추출이 비활성화됩니다.")
            self.extract_images = False
    
    def load(self, source: str) -> Document:
        """
        PDF 파일 로드
        
        Args:
            source: 파일 경로
            
        Returns:
            Document 객체
        """
        if PyPDF2 is None:
            logger.error("PyPDF2 패키지가 설치되지 않았습니다. 'pip install PyPDF2>=3.0.0'를 실행하여 설치하세요.")
            return Document(content="", metadata={"error": "PyPDF2 패키지가 설치되지 않았습니다.", "source": source})
        
        try:
            # PyPDF2로 텍스트 추출
            with open(source, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                num_pages = len(reader.pages)
                
                content = ""
                for i in range(num_pages):
                    page = reader.pages[i]
                    content += page.extract_text() + "\n\n"
            
            metadata = self._extract_metadata(source)
            metadata["num_pages"] = num_pages
            
            # 이미지 추출 (선택 사항)
            if self.extract_images and fitz is not None:
                try:
                    doc = fitz.open(source)
                    image_list = []
                    
                    for i, page in enumerate(doc):
                        images = page.get_images(full=True)
                        for img_index, img in enumerate(images):
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            image_list.append({
                                "page": i,
                                "index": img_index,
                                "width": base_image["width"],
                                "height": base_image["height"],
                                "format": base_image["ext"]
                            })
                    
                    metadata["images"] = image_list
                    metadata["has_images"] = len(image_list) > 0
                
                except Exception as e:
                    logger.warning(f"이미지 추출 중 오류 발생: {e}")
            
            return Document(content=content, metadata=metadata)
        
        except Exception as e:
            logger.error(f"PDF 파일 로드 중 오류 발생: {e}")
            return Document(content="", metadata={"error": str(e), "source": source})


class DocxLoader(DocumentLoader):
    """Word 문서 로더"""
    
    def __init__(self, **kwargs):
        """
        Word 문서 로더 초기화
        
        Args:
            **kwargs: 추가 매개변수
        """
        super().__init__(**kwargs)
        
        if docx is None:
            logger.warning("python-docx 패키지가 설치되지 않았습니다. DOCX 로드가 제한됩니다.")
    
    def load(self, source: str) -> Document:
        """
        Word 문서 로드
        
        Args:
            source: 파일 경로
            
        Returns:
            Document 객체
        """
        if docx is None:
            logger.error("python-docx 패키지가 설치되지 않았습니다. 'pip install python-docx>=0.8.11'를 실행하여 설치하세요.")
            return Document(content="", metadata={"error": "python-docx 패키지가 설치되지 않았습니다.", "source": source})
        
        try:
            # python-docx로 텍스트 추출
            doc = docx.Document(source)
            
            # 단락 추출
            paragraphs = [p.text for p in doc.paragraphs]
            
            # 표 추출
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            # 내용 결합
            content = "\n\n".join(paragraphs)
            
            # 표 내용 추가
            if tables:
                content += "\n\n표 내용:\n"
                for i, table in enumerate(tables):
                    content += f"\n표 {i+1}:\n"
                    for row in table:
                        content += " | ".join(row) + "\n"
            
            metadata = self._extract_metadata(source)
            metadata["num_paragraphs"] = len(paragraphs)
            metadata["num_tables"] = len(tables)
            
            return Document(content=content, metadata=metadata)
        
        except Exception as e:
            logger.error(f"Word 문서 로드 중 오류 발생: {e}")
            return Document(content="", metadata={"error": str(e), "source": source})


class HTMLLoader(DocumentLoader):
    """HTML 파일 로더"""
    
    def __init__(self, encoding: str = "utf-8", extract_links: bool = True, **kwargs):
        """
        HTML 파일 로더 초기화
        
        Args:
            encoding: 파일 인코딩
            extract_links: 링크 추출 여부
            **kwargs: 추가 매개변수
        """
        super().__init__(**kwargs)
        self.encoding = encoding
        self.extract_links = extract_links
        
        if BeautifulSoup is None:
            logger.warning("BeautifulSoup 패키지가 설치되지 않았습니다. HTML 로드가 제한됩니다.")
    
    def load(self, source: str) -> Document:
        """
        HTML 파일 로드
        
        Args:
            source: 파일 경로 또는 URL
            
        Returns:
            Document 객체
        """
        if BeautifulSoup is None:
            logger.error("BeautifulSoup 패키지가 설치되지 않았습니다. 'pip install beautifulsoup4>=4.10.0'를 실행하여 설치하세요.")
            return Document(content="", metadata={"error": "BeautifulSoup 패키지가 설치되지 않았습니다.", "source": source})
        
        try:
            # URL인지 확인
            is_url = source.startswith(("http://", "https://"))
            
            if is_url:
                if requests is None:
                    logger.error("requests 패키지가 설치되지 않았습니다. 'pip install requests>=2.28.0'를 실행하여 설치하세요.")
                    return Document(content="", metadata={"error": "requests 패키지가 설치되지 않았습니다.", "source": source})
                
                # URL에서 HTML 가져오기
                response = requests.get(source)
                response.raise_for_status()
                html = response.text
            else:
                # 파일에서 HTML 읽기
                with open(source, 'r', encoding=self.encoding) as f:
                    html = f.read()
            
            # BeautifulSoup으로 파싱
            soup = BeautifulSoup(html, 'html.parser')
            
            # 텍스트 추출
            content = soup.get_text(separator="\n", strip=True)
            
            metadata = self._extract_metadata(source)
            metadata["is_url"] = is_url
            
            # 링크 추출 (선택 사항)
            if self.extract_links:
                links = []
                for link in soup.find_all('a', href=True):
                    href = link['href']
                    text = link.get_text(strip=True)
                    links.append({"href": href, "text": text})
                
                metadata["links"] = links
                metadata["num_links"] = len(links)
            
            return Document(content=content, metadata=metadata)
        
        except Exception as e:
            logger.error(f"HTML 로드 중 오류 발생: {e}")
            return Document(content="", metadata={"error": str(e), "source": source})


class CSVLoader(DocumentLoader):
    """CSV 파일 로더"""
    
    def __init__(self, encoding: str = "utf-8", delimiter: str = ",", **kwargs):
        """
        CSV 파일 로더 초기화
        
        Args:
            encoding: 파일 인코딩
            delimiter: 구분자
            **kwargs: 추가 매개변수
        """
        super().__init__(**kwargs)
        self.encoding = encoding
        self.delimiter = delimiter
    
    def load(self, source: str) -> Document:
        """
        CSV 파일 로드
        
        Args:
            source: 파일 경로
            
        Returns:
            Document 객체
        """
        try:
            import csv
            
            rows = []
            with open(source, 'r', encoding=self.encoding, newline='') as f:
                reader = csv.reader(f, delimiter=self.delimiter)
                for row in reader:
                    rows.append(row)
            
            if not rows:
                return Document(content="", metadata={"error": "빈 CSV 파일", "source": source})
            
            # 헤더 추출
            headers = rows[0]
            
            # 내용 변환
            content = ""
            for i, row in enumerate(rows):
                if i == 0:
                    content += " | ".join(row) + "\n"
                    content += "-" * (sum(len(cell) for cell in row) + 3 * (len(row) - 1)) + "\n"
                else:
                    content += " | ".join(row) + "\n"
            
            metadata = self._extract_metadata(source)
            metadata["num_rows"] = len(rows)
            metadata["num_columns"] = len(headers)
            metadata["headers"] = headers
            
            return Document(content=content, metadata=metadata)
        
        except Exception as e:
            logger.error(f"CSV 파일 로드 중 오류 발생: {e}")
            return Document(content="", metadata={"error": str(e), "source": source})


class DirectoryLoader(DocumentLoader):
    """디렉토리 로더"""
    
    def __init__(self, recursive: bool = True, **kwargs):
        """
        디렉토리 로더 초기화
        
        Args:
            recursive: 재귀적 로드 여부
            **kwargs: 추가 매개변수
        """
        super().__init__(**kwargs)
        self.recursive = recursive
        
        # 파일 확장자별 로더 매핑
        self.loaders = {
            ".txt": TextLoader(**kwargs),
            ".md": MarkdownLoader(**kwargs),
            ".pdf": PDFLoader(**kwargs),
            ".docx": DocxLoader(**kwargs),
            ".html": HTMLLoader(**kwargs),
            ".htm": HTMLLoader(**kwargs),
            ".csv": CSVLoader(**kwargs)
        }
    
    def load(self, source: str) -> List[Document]:
        """
        디렉토리 로드
        
        Args:
            source: 디렉토리 경로
            
        Returns:
            Document 객체 목록
        """
        if not os.path.isdir(source):
            logger.error(f"'{source}'는 디렉토리가 아닙니다.")
            return []
        
        documents = []
        
        try:
            # 파일 목록 가져오기
            if self.recursive:
                for root, _, files in os.walk(source):
                    for file in files:
                        file_path = os.path.join(root, file)
                        doc = self._load_file(file_path)
                        if doc:
                            documents.append(doc)
            else:
                for file in os.listdir(source):
                    file_path = os.path.join(source, file)
                    if os.path.isfile(file_path):
                        doc = self._load_file(file_path)
                        if doc:
                            documents.append(doc)
            
            return documents
        
        except Exception as e:
            logger.error(f"디렉토리 로드 중 오류 발생: {e}")
            return []
    
    def _load_file(self, file_path: str) -> Optional[Document]:
        """
        파일 로드
        
        Args:
            file_path: 파일 경로
            
        Returns:
            Document 객체 또는 None
        """
        _, ext = os.path.splitext(file_path.lower())
        
        if ext in self.loaders:
            try:
                return self.loaders[ext].load(file_path)
            except Exception as e:
                logger.error(f"파일 '{file_path}' 로드 중 오류 발생: {e}")
        
        return None 